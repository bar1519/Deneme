import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import requests
from datetime import datetime
from openai import OpenAI
import google.generativeai as genai  # Google Gemini Entegrasyonu
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

st.set_page_config(
    page_title="Onay Mekanizmalı Çoklu Ajan Raporlama Sistemi",
    layout="wide",
)

st.title("Çoklu Ajan Raporlama Sistemi (Adım Adım Onaylı)")

NVIDIA_BASE_URL = "https://integrate.api.nvidia.com/v1"


def get_secret(name: str, default: str = "") -> str:
    try:
        value = st.secrets.get(name, default)
        if value is None:
            return default
        text = str(value).strip().strip('"').strip("'")
        if text.startswith("$") or text in ("", "None"):
            return default
        return text
    except Exception:
        return default


def get_secret_float(name: str, default: float) -> float:
    raw = get_secret(name, "")
    if not raw:
        return default
    try:
        return float(raw)
    except ValueError:
        return default


def get_secret_int(name: str, default: int) -> int:
    raw = get_secret(name, "")
    if not raw:
        return default
    try:
        return int(float(raw))
    except ValueError:
        return default


VISION_MODEL = get_secret("VISION_MODEL", "meta/llama-3.2-11b-vision-instruct")
VISION_API_KEY = get_secret("VISION_API_KEY") or get_secret("AGENT1_API_KEY")
VISION_MAX_IMAGES = get_secret_int("VISION_MAX_IMAGES", 8)
VISION_TEMPERATURE = get_secret_float("VISION_TEMPERATURE", 1.0)
VISION_TOP_P = get_secret_float("VISION_TOP_P", 1.0)
VISION_MAX_TOKENS = get_secret_int("VISION_MAX_TOKENS", 512)

# Her ajan: kendi API key + model + parametreler (Secrets)
# 1: Llama 3.1 70B (NVIDIA) | 2: DeepSeek V4 Pro (NVIDIA) | 3: Gemini (Google)
AGENT_CONFIG = {
    1: {
        "name": "Llama 3.1 70B (NVIDIA)",
        "api_key": get_secret("AGENT1_API_KEY"),
        "model": get_secret("AGENT1_MODEL", "meta/llama-3.1-70b-instruct"),
        "temperature": get_secret_float("AGENT1_TEMPERATURE", 0.2),
        "top_p": get_secret_float("AGENT1_TOP_P", 0.7),
        "max_tokens": get_secret_int("AGENT1_MAX_TOKENS", 4096),
        "frequency_penalty": None,
        "presence_penalty": None,
        "extra_body": None,
    },
    2: {
        "name": "DeepSeek V4 Pro (NVIDIA)",
        "api_key": get_secret("AGENT2_API_KEY"),
        "model": get_secret("AGENT2_MODEL", "deepseek-ai/deepseek-v4-pro"),
        "temperature": get_secret_float("AGENT2_TEMPERATURE", 1.0),
        "top_p": get_secret_float("AGENT2_TOP_P", 0.95),
        "max_tokens": get_secret_int("AGENT2_MAX_TOKENS", 16384),
        "frequency_penalty": None,
        "presence_penalty": None,
        "extra_body": {"chat_template_kwargs": {"thinking": False}},
    },
    3: {
        "name": "Gemini Raporlama Üstadı (Google)",
        "api_key": get_secret("GEMINI_API_KEY"),
        "model": get_secret("AGENT3_MODEL", "gemini-3.5-flash"),
        "temperature": get_secret_float("AGENT3_TEMPERATURE", 0.2),
        "top_p": get_secret_float("AGENT3_TOP_P", 1.0),
        "max_tokens": get_secret_int("AGENT3_MAX_TOKENS", 8192),
        "frequency_penalty": None,
        "presence_penalty": None,
        "extra_body": None,
    },
}

# --- YAN MENÜ ---
st.sidebar.header("Ajan / API Durumu")
st.sidebar.caption("Anahtarlar Streamlit Secrets'tan okunur.")
for no, cfg in AGENT_CONFIG.items():
    if no in [1, 2]:
        key_ok = bool(cfg["api_key"]) and cfg["api_key"].startswith("nvapi-") and len(cfg["api_key"]) > 10
    else:
        # Gemini API Key kontrolü
        key_ok = bool(cfg["api_key"]) and len(cfg["api_key"]) > 15

    st.sidebar.write(f"Ajan {no} ({cfg['name']}): {'key OK' if key_ok else 'key YOK'}")
    st.sidebar.caption(cfg["model"])

st.sidebar.write(
    f"Vision: {'key OK' if (VISION_API_KEY and VISION_API_KEY.startswith('nvapi-') and len(VISION_API_KEY) > 10) else 'key YOK'}"
)
st.sidebar.caption(VISION_MODEL)

# --- SESSION STATE ---
defaults = {
    "current_step": 1,  # 1=upload, 2=ajan1, 3=ajan2, 4=ajan3, 5=final
    "raw_data": None,
    "agent1_output": None,
    "agent2_output": None,
    "agent3_output": None,
    "final_report": None,
    "finalized_at": None,  # 1 | 2 | 3
    "chat_log": [],  # [{role, content, agent, step}]
    "user_instruction_1": "",
    "user_instruction_2": "",
    "user_instruction_3": "",
    "image_analyses": [],
}
for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

AGENT_PROMPTS = {
    1: (
        "Sen 1. Ajan (Veri Çıkarım ve Yapılandırma)sın. Görevin, sana sağlanan ham veri "
        "ve metin yığınlarını okumak, gereksiz tekrarları temizlemek ve veriyi mantıklı "
        "başlıklar altında yapılandırılmış teknik bir özet haline getirmektir. "
        "Yorum ekleme, sadece veriyi düzenle. "
        "Ham veri içinde 'GÖRSEL ANALİZİ' bölümleri varsa bunları da yapılandırmaya dahil et "
        "(tablolar, grafikler, ekran görüntüleri, belgelerdeki resimler). "
        "KRİTİK KURALLAR: (1) Yalnızca verilen ham verideki gerçek isimleri, sayıları ve "
        "değerleri kullan. (2) [Bölüm 1], [ortalama puan], örnek/şablon/yer tutucu metin "
        "UYDURMA. (3) Ham veri boşsa veya okunamadıysa bunu açıkça yaz; sahte veri üretme."
    ),
    2: (
        "Sen 2. Ajan (Teknik Analiz)sin. Sana hem ham kaynak veri hem de 1. ajanın "
        "yapılandırılmış çıktısı verilir. Eğilimleri, anormallikleri, kritik değerleri "
        "ve teknik çıkarımları belirleyerek detaylı bir analitik rapor taslağı oluştur. "
        "1. ajan çıktısını temel al, ham veriyle doğrula ve zenginleştir. "
        "Görsel analiz metinlerini (grafik, tablo görüntüsü, belge resmi) sayısal/metinsel "
        "veriyle karşılaştır; çelişki veya ek bulgu varsa belirt. "
        "Ham veri veya 1. ajan çıktısı boş/şablon ise uydurma analiz yapma; eksikliği belirt."
    ),
    3: (
        "Sen 3. Ajan (Nihai Rapor, Kurumsal Biçimlendirme ve Görselleştirme Üstadı)sın. "
        "Sana ham kaynak veri, 1. ajan çıktısı ve 2. ajan analizi verilir. "
        "Bunları birleştirerek tutarlı, eksiksiz, üst düzey yöneticilere sunuma hazır nihai teknik raporu yaz.\n"
        "KESİNLİKLE ŞU 4 KURALLA HAREKET ET:\n"
        "1. YÖNETİCİ ÖZETİ: Raporun en tepesine yöneticilerin tek bakışta konuyu, riskleri ve sonuçları anlayacağı net bir 'Yönetici Özeti' (Executive Summary) yaz.\n"
        "2. NİZAMİ TABLOLAR: Metin içindeki tüm veri karşılaştırmalarını, anomali listlerini ve sayısal dağılımları KESİNLİKLE nizami 'Markdown Tabloları' haline getir.\n"
        "3. MERMAID ŞEMALARI: Süreç akışlarını, veri ilişkilerini veya dağılımları görselleştirmek için mutlaka en az bir adet geçerli 'mermaid' kod bloğu (örneğin flowchart TD, sequenceDiagram veya pie chart) oluştur.\n"
        "4. KURUMSAL DİL: Dili aşırı övgüden arındırılmış, doğrudan, teknik odaklı ve kusursuz bir kurumsal Türkçe ile yaz. Asla yer tutucu veya örnek veri üretme."
    ),
}


def init_nvidia_client(agent_no: int = 1):
    cfg = AGENT_CONFIG[agent_no]
    api_key = cfg["api_key"]
    if not api_key or not api_key.startswith("nvapi-"):
        st.error(
            f"Ajan {agent_no} için geçerli API anahtarı yok. "
            f"Secrets'a AGENT{agent_no}_API_KEY = \"nvapi-...\" ekleyin."
        )
        return None
    return OpenAI(base_url=NVIDIA_BASE_URL, api_key=api_key, timeout=300.0)


def init_gemini_client():
    cfg = AGENT_CONFIG[3]
    api_key = cfg["api_key"]
    if not api_key or len(api_key) < 15:
        st.error(
            "3. Ajan (Gemini) için geçerli API anahtarı yok. "
            "Secrets'a GEMINI_API_KEY = \"AIzaSy...\" ekleyin."
        )
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(cfg["model"])


def parse_excel(file):
    """Tüm sheet'leri okur."""
    file.seek(0)
    sheets = pd.read_excel(file, sheet_name=None)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"[Sayfa: {sheet_name}]")
        if df.empty:
            parts.append("(boş sayfa)")
        else:
            parts.append(df.to_string(index=False))
    return "\n".join(parts)


def parse_docx(file):
    """Paragraflar + tablolar (çoğu puan/liste Word tablosunda olur)."""
    file.seek(0)
    doc = Document(io.BytesIO(file.read()))
    parts = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for t_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"\n--- Tablo {t_idx} ---")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    (p.text or "").strip() for p in cell.paragraphs
                ).strip()
                cells.append(cell_text.replace("\n", " "))
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _guess_mime(name: str, content_type: str = "") -> str:
    ct = (content_type or "").lower()
    if ct.startswith("image/"):
        return ct
    lower = (name or "").lower()
    if lower.endswith(".png"):
        return "image/png"
    if lower.endswith(".webp"):
        return "image/webp"
    if lower.endswith(".gif"):
        return "image/gif"
    return "image/jpeg"


def extract_images_from_docx(file):
    """Word içindeki gömülü görselleri (bytes, mime, label) listesi olarak döner."""
    file.seek(0)
    doc = Document(io.BytesIO(file.read()))
    images = []
    idx = 0
    for rel in doc.part.rels.values():
        if "image" not in (rel.reltype or ""):
            continue
        idx += 1
        try:
            blob = rel.target_part.blob
            mime = _guess_mime(
                getattr(rel.target_part, "partname", "") or "",
                getattr(rel.target_part, "content_type", "") or "",
            )
            images.append(
                {
                    "blob": blob,
                    "mime": mime,
                    "label": f"{getattr(file, 'name', 'docx')}_gorsel_{idx}",
                }
            )
        except Exception:
            continue
    return images


def collect_uploaded_image(file):
    file.seek(0)
    blob = file.read()
    mime = _guess_mime(file.name, getattr(file, "type", "") or "")
    return {"blob": blob, "mime": mime, "label": file.name}


def init_vision_client():
    """Vision için API key kontrolü (çağrı requests ile yapılır)."""
    if not VISION_API_KEY or not VISION_API_KEY.startswith("nvapi-"):
        return None
    return True


def analyze_image_with_vision(_client, image: dict) -> str:
    """NVIDIA vision: requests + chat/completions."""
    b64 = base64.b64encode(image["blob"]).decode("utf-8")
    data_url = f"data:{image['mime']};base64,{b64}"
    prompt = (
        "Bu görseli teknik rapor için incele. "
        "Varsa tüm okunabilir metinleri (OCR), tablo/grafik değerlerini, "
        "başlıkları ve dikkat çeken bulguları Türkçe, maddeler halinde yaz. "
        "Uydurma veri ekleme; okuyamadığın yerleri belirt."
    )

    invoke_url = f"{NVIDIA_BASE_URL}/chat/completions"
    stream = False
    headers = {
        "Authorization": f"Bearer {VISION_API_KEY}",
        "Accept": "text/event-stream" if stream else "application/json",
        "Content-Type": "application/json",
    }
    payload = {
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        "model": VISION_MODEL,
        "frequency_penalty": 0,
        "max_tokens": VISION_MAX_TOKENS,
        "presence_penalty": 0,
        "stream": stream,
        "temperature": VISION_TEMPERATURE,
        "top_p": VISION_TOP_P,
    }

    response = requests.post(
        invoke_url, headers=headers, json=payload, stream=stream, timeout=300
    )
    if response.status_code >= 400:
        raise RuntimeError(f"{response.status_code}: {response.text[:500]}")

    data = response.json()
    return (
        data.get("choices", [{}])[0]
        .get("message", {})
        .get("content")
        or ""
    )


def analyze_all_images(images: list) -> str:
    """Görselleri sırayla VLM ile analiz eder; metin bloğu döner."""
    if not images:
        return ""

    client = init_vision_client()
    if client is None:
        return (
            "\n=== GÖRSEL ANALİZİ ===\n"
            f"{len(images)} görsel bulundu ancak VISION_API_KEY / AGENT1_API_KEY yok; "
            "analiz atlandı.\n"
        )

    limited = images[:VISION_MAX_IMAGES]
    parts = [
        "\n=== GÖRSEL ANALİZİ ===",
        f"Toplam {len(images)} görsel; analiz edilen: {len(limited)} "
        f"(model: {VISION_MODEL})",
    ]
    for i, img in enumerate(limited, start=1):
        label = img.get("label", f"gorsel_{i}")
        try:
            if len(img["blob"]) > 4_000_000:
                parts.append(f"\n--- {label} ---\n(çok büyük, atlandı)")
                continue
            desc = analyze_image_with_vision(client, img)
            parts.append(f"\n--- {label} ---\n{desc}")
        except Exception as e:
            parts.append(f"\n--- {label} ---\nGörsel analiz hatası: {e}")

    if len(images) > VISION_MAX_IMAGES:
        parts.append(
            f"\n(Not: {len(images) - VISION_MAX_IMAGES} görsel limit nedeniyle atlandı.)"
        )
    return "\n".join(parts)


def validate_extracted_text(name: str, content: str) -> str:
    """Boş/çok kısa içerikte uyarı; ajanın uydurma şablon üretmesini engeller."""
    text = (content or "").strip()
    if len(text) < 20:
        return (
            f"{name}: DOSYA İÇERİĞİ BOŞ VEYA OKUNAMADI. "
            "Paragraf/tablo bulunamadı. Lütfen dosyayı kontrol edin."
        )
    return text


def call_agent(client, agent_no: int, user_content: str, max_tokens=None):
    cfg = AGENT_CONFIG[agent_no]
    instruction = (st.session_state.get(f"user_instruction_{agent_no}") or "").strip()
    if instruction:
        user_content = (
            f"=== KULLANICI TALİMATI (bu adımda yapılacaklar) ===\n"
            f"{instruction}\n\n"
            f"{user_content}"
        )

    # 1. ve 2. Ajan NVIDIA (OpenAI Client) Kullanır
    if agent_no in [1, 2]:
        kwargs = {
            "model": cfg["model"],
            "messages": [
                {"role": "system", "content": AGENT_PROMPTS[agent_no]},
                {"role": "user", "content": user_content},
            ],
            "temperature": cfg["temperature"],
            "top_p": cfg["top_p"],
            "max_tokens": cfg["max_tokens"] if max_tokens is None else max_tokens,
            "stream": False,
        }
        if cfg.get("frequency_penalty") is not None:
            kwargs["frequency_penalty"] = cfg["frequency_penalty"]
        if cfg.get("presence_penalty") is not None:
            kwargs["presence_penalty"] = cfg["presence_penalty"]
        if cfg.get("extra_body"):
            kwargs["extra_body"] = cfg["extra_body"]

        completion = client.chat.completions.create(**kwargs)
        content = completion.choices[0].message.content
        return content if content is not None else ""

    # 3. Ajan Google Gemini Kullanır
    elif agent_no == 3:
        prompt_with_system = f"{AGENT_PROMPTS[3]}\n\n{user_content}"
        generation_config = genai.types.GenerationConfig(
            temperature=cfg["temperature"],
            top_p=cfg["top_p"],
            max_output_tokens=cfg["max_tokens"] if max_tokens is None else max_tokens,
        )
        response = client.generate_content(
            prompt_with_system,
            generation_config=generation_config
        )
        return response.text if response.text else ""


def append_chat(role: str, content: str, agent=None):
    st.session_state.chat_log.append(
        {
            "role": role,
            "content": content,
            "agent": agent,
            "step": st.session_state.current_step,
        }
    )


def render_chat_panel(agent_no: int, placeholder: str):
    """Her adımda kullanıcının talimat yazdığı chat alanı."""
    st.subheader("Chat / Talimat")
    st.caption(
        f"Ajan {agent_no} için bu adımda ne yapılsın? Yazıların ajan çağrısına eklenir."
    )

    relevant = [
        m
        for m in st.session_state.chat_log
        if m.get("agent") in (None, agent_no)
    ]
    if relevant:
        with st.container(height=220):
            for m in relevant[-30:]:
                who = {
                    "user": "Sen",
                    "assistant": f"Ajan {m.get('agent') or '?'}",
                    "system": "Sistem",
                }.get(m["role"], m["role"])
                st.markdown(f"**{who}:** {m['content']}")
    else:
        st.caption("Henüz mesaj yok. Aşağıya bu adım için talimatını yaz.")

    note_key = f"user_instruction_{agent_no}"
    st.text_area(
        "Bu adımda yapılacaklar (sen yaz)",
        height=120,
        key=note_key,
        placeholder=placeholder,
    )

    if st.button("Talimatı sohbete kaydet", key=f"save_chat_{agent_no}"):
        text = (st.session_state.get(note_key) or "").strip()
        if text:
            append_chat("user", text, agent=agent_no)
            st.rerun()


def build_agent2_prompt(raw_data: str, agent1_output: str) -> str:
    return (
        "=== HAM KAYNAK VERİ (1. ajana giden aynı veri) ===\n"
        f"{raw_data}\n\n"
        "=== 1. AJAN ÇIKTISI ===\n"
        f"{agent1_output}\n\n"
        "Yukarıdaki ham veri ile 1. ajan çıktısını birlikte kullanarak teknik analiz yap."
    )


def build_agent3_prompt(raw_data: str, agent1_output: str, agent2_output: str) -> str:
    return (
        "=== HAM KAYNAK VERİ (ilk veri) ===\n"
        f"{raw_data}\n\n"
        "=== 1. AJAN ÇIKTISI ===\n"
        f"{agent1_output}\n\n"
        "=== 2. AJAN ÇIKTISI ===\n"
        f"{agent2_output}\n\n"
        "Ham veri + 1. ajan + 2. ajan çıktılarını birleştirerek Yönetici Özeti, "
        "Markdown tabloları ve Mermaid şeması içeren kusursuz nihai raporu oluştur."
    )


def reset_pipeline():
    st.session_state.current_step = 1
    st.session_state.raw_data = None
    st.session_state.agent1_output = None
    st.session_state.agent2_output = None
    st.session_state.agent3_output = None
    st.session_state.final_report = None
    st.session_state.finalized_at = None
    st.session_state.chat_log = []
    st.session_state.user_instruction_1 = ""
    st.session_state.user_instruction_2 = ""
    st.session_state.user_instruction_3 = ""
    st.session_state.image_analyses = []


def finalize_with(output_text: str, agent_no: int):
    st.session_state.final_report = output_text
    st.session_state.finalized_at = agent_no
    st.session_state.current_step = 5
    st.rerun()


def _report_title(agent_no: int) -> str:
    return f"Teknik Rapor (Ajan {agent_no} çıktısı)"


def build_docx_bytes(report_text: str, agent_no: int) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(_report_title(agent_no), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    for line in (report_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _register_pdf_font() -> str:
    candidates = [
        (r"C:\Windows\Fonts\arial.ttf", "ArialTR"),
        (r"C:\Windows\Fonts\calibri.ttf", "CalibriTR"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVu"),
    ]
    for path, name in candidates:
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return name
        except Exception:
            continue
    return "Helvetica"


def build_pdf_bytes(report_text: str, agent_no: int) -> bytes:
    font_name = _register_pdf_font()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        spaceAfter=8,
        alignment=1,
    )
    meta_style = ParagraphStyle(
        "ReportMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        spaceAfter=16,
        alignment=1,
    )
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=15,
        spaceAfter=6,
    )
    h1_style = ParagraphStyle(
        "ReportH1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "ReportH2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12,
        spaceBefore=10,
        spaceAfter=4,
    )

    def esc(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    story = [
        Paragraph(esc(_report_title(agent_no)), title_style),
        Paragraph(
            esc(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}"),
            meta_style,
        ),
        Spacer(1, 0.3 * cm),
    ]

    for line in (report_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.25 * cm))
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(esc(stripped[4:]), h2_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(esc(stripped[3:]), h2_style))
        elif stripped.startswith("# "):
            story.append(Paragraph(esc(stripped[2:]), h1_style))
        else:
            story.append(Paragraph(esc(stripped), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def show_progress():
    labels = ["Yükleme", "1. Ajan", "2. Ajan", "3. Ajan", "Sonuç"]
    step = st.session_state.current_step
    cols = st.columns(5)
    for i, (col, label) in enumerate(zip(cols, labels), start=1):
        with col:
            if i < step:
                st.success(label)
            elif i == step:
                st.info(f"→ {label}")
            else:
                st.caption(label)


show_progress()

# --- ADIM 1: DOSYA YÜKLEME + CHAT ---
if st.session_state.current_step == 1:
    st.header("Adım 1: Kaynak Dosyaların Yüklenmesi")
    st.caption(
        "Ham veri önce 1. ajana gider. Solda dosya yükle, sağda bu adımda ne yapılacağını yaz."
    )

    left, right = st.columns([1.2, 1])
    with left:
        uploaded_files = st.file_uploader(
            "Excel / Word / görsel dosyaları seçin",
            type=["xlsx", "docx", "png", "jpg", "jpeg", "webp"],
            accept_multiple_files=True,
            help="Word içindeki gömülü resimler de otomatik incelenir.",
        )
    with right:
        render_chat_panel(
            1,
            "Örn: Tekrarları sil, tabloları düzenle, görsellerdeki puanları da dahil et...",
        )

    if uploaded_files and st.button("Dosyaları İşle ve 1. Ajanı Çalıştır"):
        client = init_nvidia_client(1)
        if client:
            with st.spinner(
                "Dosyalar ve görseller okunuyor, vision + 1. ajan çalışıyor..."
            ):
                combined = []
                all_images = []
                for file in uploaded_files:
                    name = file.name.lower()
                    if name.endswith(".xlsx"):
                        content = validate_extracted_text(
                            file.name, parse_excel(file)
                        )
                        combined.append(f"\n--- {file.name} (Excel) ---\n{content}")
                    elif name.endswith(".docx"):
                        content = validate_extracted_text(
                            file.name, parse_docx(file)
                        )
                        combined.append(f"\n--- {file.name} (Word) ---\n{content}")
                        all_images.extend(extract_images_from_docx(file))
                    elif name.endswith((".png", ".jpg", ".jpeg", ".webp")):
                        all_images.append(collect_uploaded_image(file))
                        combined.append(
                            f"\n--- {file.name} (Görsel dosya) ---\n"
                            "(içerik vision modeliyle aşağıda analiz edilecek)"
                        )

                vision_text = ""
                if all_images:
                    with st.status(
                        f"{len(all_images)} görsel inceleniyor...", expanded=True
                    ) as status:
                        vision_text = analyze_all_images(all_images)
                        st.session_state.image_analyses = [
                            {"label": img["label"], "mime": img["mime"]}
                            for img in all_images
                        ]
                        status.update(
                            label=f"{min(len(all_images), VISION_MAX_IMAGES)} görsel analiz edildi",
                            state="complete",
                        )
                    if vision_text:
                        combined.append(vision_text)

                raw_data = "\n".join(combined).strip()
                st.session_state.raw_data = raw_data

                only_empty_docs = (
                    "DOSYA İÇERİĞİ BOŞ" in raw_data
                    and "=== GÖRSEL ANALİZİ ===" not in raw_data
                )
                if only_empty_docs or len(raw_data) < 40:
                    st.error(
                        "Dosya/görsel içeriği okunamadı veya boş. "
                        "Word tabloları ve gömülü resimler desteklenir; "
                        "ayrıca png/jpg yükleyebilirsiniz."
                    )
                    st.code(raw_data[:2000] if raw_data else "(boş)")
                    st.stop()

                instr = (st.session_state.user_instruction_1 or "").strip()
                if instr:
                    append_chat("user", instr, agent=1)

                try:
                    output = call_agent(
                        client,
                        1,
                        f"Yapılandırılacak Ham Veri:\n{raw_data}",
                    )
                    st.session_state.agent1_output = output
                    append_chat(
                        "assistant",
                        output[:1500] + ("..." if len(output) > 1500 else ""),
                        agent=1,
                    )
                    st.session_state.current_step = 2
                    st.rerun()
                except Exception as e:
                    st.error(f"API çağrısı sırasında hata: {e}")

# --- ADIM 2: 1. AJAN ÇIKTISI + CHAT (2. ajan talimatı) ---
elif st.session_state.current_step == 2:
    st.header("Adım 2: 1. Ajan Çıktısı")
    st.info(
        "1. ajan çıktısını kontrol et. 2. ajana gitmeden önce sağdaki chate "
        "analiz talimatını yazabilirsin."
    )

    left, right = st.columns([1.2, 1])
    with left:
        with st.expander("Ham kaynak veri (önizleme)", expanded=False):
            st.text(
                st.session_state.raw_data[:4000]
                + ("..." if len(st.session_state.raw_data or "") > 4000 else "")
            )

        edited_a1 = st.text_area(
            "1. Ajan çıktısı (düzenlenebilir)",
            value=st.session_state.agent1_output or "",
            height=400,
            key="edit_agent1",
        )
    with right:
        render_chat_panel(
            2,
            "Örn: Anomalilere odaklan, aylık trendleri çıkar, kritik eşikleri işaretle... (DeepSeek)",
        )

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("Baştan başla"):
            reset_pipeline()
            st.rerun()
    with c2:
        if st.button("Bu çıktıyla sonuçlandır", type="primary"):
            st.session_state.agent1_output = edited_a1
            finalize_with(edited_a1, 1)
    with c3:
        if st.button("Onayla → 2. Ajana gönder"):
            st.session_state.agent1_output = edited_a1
            client = init_nvidia_client(2)
            if client:
                with st.spinner("2. ajan analiz ediyor (ham veri + 1. ajan çıktısı)..."):
                    instr = (st.session_state.user_instruction_2 or "").strip()
                    if instr:
                        append_chat("user", instr, agent=2)
                    try:
                        prompt = build_agent2_prompt(
                            st.session_state.raw_data,
                            st.session_state.agent1_output,
                        )
                        output = call_agent(client, 2, prompt)
                        st.session_state.agent2_output = output
                        append_chat(
                            "assistant",
                            output[:1500] + ("..." if len(output) > 1500 else ""),
                            agent=2,
                        )
                        st.session_state.current_step = 3
                        st.rerun()
                    except Exception as e:
                        st.error(f"API çağrısı sırasında hata: {e}")

# --- ADIM 3: 2. AJAN ÇIKTISI + CHAT (3. ajan talimatı) ---
elif st.session_state.current_step == 3:
    st.header("Adım 3: 2. Ajan Çıktısı")
    st.info(
        "2. ajan çıktısını kontrol et. 3. ajana (Gemini) gitmeden önce nihai rapor talimatını yaz."
    )

    left, right = st.columns([1.2, 1])
    with left:
        with st.expander("Girdi özeti", expanded=False):
            st.markdown("**Ham veri** (kısaltılmış)")
            st.text((st.session_state.raw_data or "")[:2000])
            st.markdown("**1. ajan çıktısı**")
            st.text(st.session_state.agent1_output or "")

        edited_a2 = st.text_area(
            "2. Ajan çıktısı (düzenlenebilir)",
            value=st.session_state.agent2_output or "",
            height=400,
            key="edit_agent2",
        )
    with right:
        render_chat_panel(
            3,
            "Örn: Yönetici özeti yaz, riskleri maddele, net tablolar kullan... (Gemini)",
        )

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("← 1. ajan çıktısına dön"):
            st.session_state.current_step = 2
            st.rerun()
    with c2:
        if st.button("Bu çıktıyla sonuçlandır", type="primary"):
            st.session_state.agent2_output = edited_a2
            finalize_with(edited_a2, 2)
    with c3:
        if st.button("Onayla → 3. Ajana (Gemini) gönder"):
            st.session_state.agent2_output = edited_a2
            client = init_gemini_client()
            if client:
                with st.spinner(
                    "3. Ajan (Gemini) tablo ve şemalarla nihai raporu yazıyor..."
                ):
                    instr = (st.session_state.user_instruction_3 or "").strip()
                    if instr:
                        append_chat("user", instr, agent=3)
                    try:
                        prompt = build_agent3_prompt(
                            st.session_state.raw_data,
                            st.session_state.agent1_output,
                            st.session_state.agent2_output,
                        )
                        output = call_agent(client, 3, prompt)
                        st.session_state.agent3_output = output
                        append_chat(
                            "assistant",
                            output[:1500] + ("..." if len(output) > 1500 else ""),
                            agent=3,
                        )
                        st.session_state.current_step = 4
                        st.rerun()
                    except Exception as e:
                        st.error(f"Gemini API çağrısı sırasında hata: {e}")

# --- ADIM 4: 3. AJAN ÇIKTISI ---
elif st.session_state.current_step == 4:
    st.header("Adım 4: 3. Ajan (Gemini) Çıktısı")
    st.warning(
        "3. ajan ham veri + 1. ajan + 2. ajan çıktılarını birleştirdi. "
        "Düzenleyip sonuçlandırın."
    )

    with st.expander("Girdi özeti", expanded=False):
        st.markdown("**Ham veri** (kısaltılmış)")
        st.text((st.session_state.raw_data or "")[:1500])
        st.markdown("**1. ajan**")
        st.text(st.session_state.agent1_output or "")
        st.markdown("**2. ajan**")
        st.text(st.session_state.agent2_output or "")

    edited_a3 = st.text_area(
        "3. Ajan çıktısı (düzenlenebilir)",
        value=st.session_state.agent3_output or "",
        height=400,
        key="edit_agent3",
    )

    c1, c2 = st.columns(2)
    with c1:
        if st.button("← 2. ajan çıktısına dön"):
            st.session_state.current_step = 3
            st.rerun()
    with c2:
        if st.button("Nihai rapor olarak sonuçlandır", type="primary"):
            st.session_state.agent3_output = edited_a3
            finalize_with(edited_a3, 3)

# --- ADIM 5: SONUÇ ---
elif st.session_state.current_step == 5:
    st.header("Nihai Sonuç")
    finished = st.session_state.finalized_at
    st.success(f"Rapor {finished}. ajan çıktısı üzerinden sonuçlandırıldı.")

    report_text = st.session_state.final_report or ""
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    base_name = f"rapor_ajan{finished}_{stamp}"

    st.text_area(
        "Nihai rapor (Düzenleme Ekranı)",
        value=report_text,
        height=300,
        key="final_view",
    )
    
    st.subheader("Görsel Rapor Önizlemesi (Tablo ve Şema Desteği ile)")
    st.markdown(report_text, unsafe_allow_html=True)

    st.subheader("İndir")
    d1, d2, d3 = st.columns(3)
    with d1:
        st.download_button(
            "TXT indir",
            data=report_text,
            file_name=f"{base_name}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with d2:
        st.download_button(
            "Word (.docx) indir",
            data=build_docx_bytes(report_text, finished),
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with d3:
        st.download_button(
            "PDF indir",
            data=build_pdf_bytes(report_text, finished),
            file_name=f"{base_name}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

    if st.button("Yeni analiz başlat"):
        reset_pipeline()
        st.rerun()
